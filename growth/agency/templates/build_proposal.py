#!/usr/bin/env python3
"""
Build the tsukumo consulting proposal / SOW template as a .docx.

Usage (needs python-docx):
    python3 -m venv .venv && .venv/bin/pip install python-docx
    .venv/bin/python build_proposal.py        # writes tsukumo-proposal-template.docx

Brand: tsukumo only (no Synergix in prose). Acid-lime accent #c8ff00 on a near-black base.
No fabricated numbers: pricing + any metrics are bracketed placeholders for the owner to fill.
Proof is qualitative (engagement summaries) until verified numbers land.
Voice: prod-grade, builder-credible, dev-respecting, augment-not-replace.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACID = RGBColor(0xC8, 0xFF, 0x00)
INK = RGBColor(0x0B, 0x0E, 0x14)
GREY = RGBColor(0x6B, 0x72, 0x80)

DOC = Document()

# Base styles
normal = DOC.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK


def heading(text, size=16, color=INK, space_before=14, space_after=6, caps=False):
    p = DOC.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text.upper() if caps else text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def body(text, color=INK, size=10.5, italic=False):
    p = DOC.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.italic = italic
    return p


def bullet(text):
    p = DOC.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10.5)
    return p


def placeholder(text):
    p = DOC.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.bold = True
    r.font.color.rgb = RGBColor(0x99, 0x66, 0x00)  # amber: fill before sending
    return p


# ── Cover ────────────────────────────────────────────────────────────────────
title = DOC.add_paragraph()
title.paragraph_format.space_before = Pt(40)
tr = title.add_run("tsukumo")
tr.bold = True
tr.font.size = Pt(34)
tr.font.color.rgb = INK
sub = DOC.add_paragraph()
sr = sub.add_run("we run AI in production")
sr.font.size = Pt(13)
sr.font.color.rgb = GREY
bar = DOC.add_paragraph()
br = bar.add_run("—" * 12)
br.font.color.rgb = ACID
br.bold = True

heading("Proposal & Statement of Work", size=20)
placeholder("[CLIENT NAME] · [DATE] · prepared for [BUYER NAME, TITLE]")
body("Confidential. Prepared by tsukumo for the named recipient.", color=GREY, italic=True)

DOC.add_page_break()

# ── 1. The problem ─────────────────────────────────────────────────────────────
heading("1. The problem", caps=True)
body("Your team has AI as a copilot. It autocompletes and answers questions in the editor, "
     "and that is roughly 10% of what coding agents can do. The other 90% is agents running "
     "real work in production: building, testing, fixing, and coordinating, operated by your "
     "developers rather than typed by them.")
body("The distance between those two is an operating problem, not a license problem. Closing "
     "it needs the right context for agents, orchestration so a fleet doesn't collide, "
     "observability so you run on evidence, and a team trained to operate, not just prompt.")
placeholder("[CLIENT-SPECIFIC: 1-2 sentences on where THIS team is stuck today — from the "
            "discovery call. Keep it concrete and true.]")

# ── 2. Our approach ────────────────────────────────────────────────────────────
heading("2. Our approach", caps=True)
body("We turn your dev team into agentic operators, on your existing environment, your "
     "standards, and your stack. We augment your developers; we do not replace them. The "
     "craft stays theirs, and the output gets much bigger.")
heading("How we work", size=12, space_before=8)
bullet("Production from the start — agents doing real, measured work early, not a POC.")
bullet("On your standards — your CI, review culture, and constraints hold.")
bullet("Augment, never replace — your devs become the operators; the capability stays.")
bullet("We run what we sell — we operate our own agent fleets to ship our own software.")

# ── 3. Scope ───────────────────────────────────────────────────────────────────
heading("3. Scope of work", caps=True)
body("Three phases, adapted to the engagement:")
heading("Phase 1 — Strategy", size=12, space_before=8)
bullet("Assess where the team actually is (not where the demos say).")
bullet("Identify the highest-impact agent workflows for your codebase.")
bullet("Set a realistic path from copilot to operator; honest about what AI won't fix.")
heading("Phase 2 — Implementation", size=12, space_before=8)
bullet("Build the agent workflows + supporting layers (context, orchestration, observability) "
       "on your repo, in production, with your controls intact.")
heading("Phase 3 — Training", size=12, space_before=8)
bullet("Bring your developers up to operator level so the capability stays after we leave.")
placeholder("[CLIENT-SPECIFIC: in/out-of-scope list agreed in scoping. Be explicit about "
            "what is NOT included.]")

# ── 4. Timeline ────────────────────────────────────────────────────────────────
heading("4. Timeline", caps=True)
placeholder("[REAL NUMBERS NEEDED: phase durations + milestones, e.g. Strategy [X weeks], "
            "Implementation [X weeks], Training [X weeks]. Set with the client.]")

# ── 5. Investment ──────────────────────────────────────────────────────────────
heading("5. Investment", caps=True)
body("Pricing reflects prod-grade capability transferred to your team, not seats or hours of "
     "demos. This is about prod-grade output and roughly 10x from the team you already trust.")
placeholder("[REAL NUMBERS NEEDED: pricing model + figures — fixed-scope per phase, "
            "retainer, or blended. Do NOT send with placeholders. Owner sets the numbers.]")

# ── 6. About tsukumo & proof ───────────────────────────────────────────────────
heading("6. About tsukumo", caps=True)
body("tsukumo is a developer studio and AI consultancy. We build our own products by running "
     "agent fleets in production (the open suite WRAI.TH, trovex, and yoru is the proof), and "
     "we transition client teams to do the same. We are developers; we understand developers "
     "and bring them across without sidelining them.")
heading("Selected engagements", size=12, space_before=8)
body("Our own studio — we ship and maintain a multi-tool open suite as a small team because "
     "the agent fleet carries the repeated build, test, documentation, and coordination work. "
     "The operating model we sell is the one we run on ourselves.")
body("A quantitative fund — AI brought into a high-stakes finance environment where "
     "correctness and risk control are non-negotiable. (Anonymized under NDA.)")
body("A Swiss real-estate company — AI applied to content and marketing at scale. "
     "(Attribution on request.)")
placeholder("[REAL NUMBERS NEEDED: add one verified outcome per engagement before sending. "
            "Until then keep these qualitative — no invented metrics, logos, or quotes.]")

# ── 7. Next steps ──────────────────────────────────────────────────────────────
heading("7. Next steps", caps=True)
bullet("Confirm scope and timeline.")
bullet("Sign this SOW.")
bullet("Kickoff: access, environment walkthrough, and the first workflow we'll put in prod.")
foot = DOC.add_paragraph()
foot.paragraph_format.space_before = Pt(16)
fr = foot.add_run("tsukumo · tsukumo.ch · we run AI in production")
fr.font.size = Pt(9)
fr.font.color.rgb = GREY

DOC.save("tsukumo-proposal-template.docx")
print("wrote tsukumo-proposal-template.docx")
